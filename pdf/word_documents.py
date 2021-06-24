from docx import Document
doc = Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
for i in range(len(doc.paragraphs[1].runs)):
    print(str(i) + 'º run ' + doc.paragraphs[1].runs[i].text)


def get_fulltext(filename):
    """
    This function gets a fulltext from a docx file, based on it's filename
    :param filename: The filename that you want to get the fulltext
    :return: The fulltext
    """
    doc = Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n\n'.join(fulltext)


print(get_fulltext('demo.docx'))
title = doc.paragraphs[0]
print(title.style)
title.style = 'Heading 1'
paragraph_1 = doc.paragraphs[1]
paragraph_1.runs[0].underline = True
paragraph_1.runs[2].small_capitals = True
doc.save('restyledfile.docx')
