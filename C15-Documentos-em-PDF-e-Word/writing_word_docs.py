import docx

doc = docx.Document()
doc.add_paragraph('Hello World!')
for i in range(5):
    doc.add_heading('Header %s' % i, i)
#for i in range(200):
 #   doc.add_paragraph('Hello Word!')
# You can add line breaks
doc.paragraphs[3].runs[0].add_break()
# ... and page breaks too:
doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
# And finally some images:
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('MyFirstWordWithPython.docx')




