import csv
import docx

file = open('names.csv')
reader = csv.reader(file)
data = list(reader)
word = docx.Document()
for i in range(len(data)):
    word.add_paragraph('Melado Productions tem o prazer de convidar')
    word.add_paragraph(data[i])
    word.add_paragraph('para o 1º Festival Melado de Filmes')


word.save('guests.docx')